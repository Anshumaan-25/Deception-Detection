#!/usr/bin/env python3
"""
SPOVNOB System Master Manual — DOCX Generator
==============================================
Generates a beautifully formatted, visually rich Word document from the
SPOVNOB_SYSTEM_MASTER_MANUAL.md, including:
  - Professional cover page
  - Chapter cover dividers
  - Color-coded callout boxes
  - Hardware topology flowcharts (built as styled tables)
  - Phase transition matrix tables
  - Tensor shape progression diagrams
  - Code blocks with monospace dark styling
  - Full appendix runbook tables
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re

# ─── COLOR PALETTE ───────────────────────────────────────────────────────────
C_BG_DARK        = RGBColor(0x0F, 0x17, 0x2A)   # Deep navy background
C_ACCENT_BLUE    = RGBColor(0x38, 0x82, 0xF6)   # Blue-500
C_ACCENT_INDIGO  = RGBColor(0x63, 0x66, 0xF1)   # Indigo-500
C_ACCENT_GREEN   = RGBColor(0x10, 0xB9, 0x81)   # Emerald-500
C_ACCENT_AMBER   = RGBColor(0xF5, 0x9E, 0x0B)   # Amber-500
C_ACCENT_RED     = RGBColor(0xEF, 0x44, 0x44)   # Red-500
C_ACCENT_PURPLE  = RGBColor(0x8B, 0x5C, 0xF6)   # Violet-500
C_WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY_100       = RGBColor(0xF3, 0xF4, 0xF6)
C_GRAY_200       = RGBColor(0xE5, 0xE7, 0xEB)
C_GRAY_400       = RGBColor(0x9C, 0xA3, 0xAF)
C_GRAY_700       = RGBColor(0x37, 0x41, 0x51)
C_GRAY_800       = RGBColor(0x1F, 0x29, 0x37)
C_GRAY_900       = RGBColor(0x11, 0x18, 0x27)
C_CODE_BG        = RGBColor(0x1E, 0x29, 0x3B)   # Code block bg
C_CODE_TEXT      = RGBColor(0x7D, 0xD3, 0xFC)   # Light blue code text
C_CHAPTER_HEADER = RGBColor(0x1E, 0x40, 0xAF)   # Blue-800

# ─── HELPERS ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color: RGBColor):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{color.rgb:06X}')
    tcPr.append(shd)

def set_cell_borders(cell, color='auto', size=4):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color if isinstance(color, str) else f'{color.rgb:06X}')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def para_format(para, left=0, space_before=0, space_after=0, line_rule=None, line_spacing=None):
    pf = para.paragraph_format
    pf.left_indent = Pt(left)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_rule:
        pf.line_spacing_rule = line_rule
    if line_spacing:
        pf.line_spacing = line_spacing

def set_run_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type_page())

def docx_break_type_page():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br

def add_pagebreak(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.runs[0] if para.runs else para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_table_width(table, width_inches):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

# ─── COVER PAGE ──────────────────────────────────────────────────────────────

def build_cover(doc):
    # Full-width cover table
    tbl = doc.add_table(rows=1, cols=1)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, C_BG_DARK)
    cell.width = Inches(6.5)

    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(40)
    cp.paragraph_format.space_after = Pt(4)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Badge
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge.paragraph_format.space_before = Pt(6)
    badge.paragraph_format.space_after = Pt(6)
    # We'll add content to the cell paragraphs directly

    cell.paragraphs[0].clear()

    # Accent stripe
    stripe_p = OxmlElement('w:p')
    stripe_pPr = OxmlElement('w:pPr')
    stripe_jc = OxmlElement('w:jc')
    stripe_jc.set(qn('w:val'), 'center')
    stripe_pPr.append(stripe_jc)
    stripe_p.append(stripe_pPr)
    stripe_r = OxmlElement('w:r')
    stripe_rPr = OxmlElement('w:rPr')
    stripe_sz = OxmlElement('w:sz')
    stripe_sz.set(qn('w:val'), '8')
    stripe_rPr.append(stripe_sz)
    stripe_r.append(stripe_rPr)
    stripe_t = OxmlElement('w:t')
    stripe_t.text = '─' * 60
    stripe_r.append(stripe_t)
    stripe_p.append(stripe_r)

    def add_cover_para(text, size=11, bold=False, color=None, space_before=6, space_after=4, align='center'):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = 'Calibri Light'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color or C_WHITE
        return p

    add_cover_para('DEFENSE-GRADE FORENSIC ARCHITECTURE', 10, False, C_ACCENT_BLUE, 30, 2)
    add_cover_para('SPOVNOB', 52, True, C_WHITE, 4, 2)
    add_cover_para('SYSTEM MASTER MANUAL', 22, False, C_GRAY_200, 0, 8)

    # Divider line
    div_p = cell.add_paragraph()
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_p.paragraph_format.space_before = Pt(4)
    div_p.paragraph_format.space_after = Pt(14)
    div_run = div_p.add_run('━' * 50)
    div_run.font.color.rgb = C_ACCENT_INDIGO
    div_run.font.size = Pt(10)

    add_cover_para('Temporal Behavioral Anomaly Detection', 14, False, C_GRAY_200, 2, 2)
    add_cover_para('& Cognitive Load Analysis Platform', 14, False, C_GRAY_200, 0, 10)

    # Key specs badges row
    specs_p = cell.add_paragraph()
    specs_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    specs_p.paragraph_format.space_before = Pt(8)
    specs_p.paragraph_format.space_after = Pt(4)
    for i, spec in enumerate(['44-Core CPU', '512 GB ECC RAM', 'RTX 6000 Ada · 48GB VRAM']):
        r = specs_p.add_run(f'  {spec}  ')
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = C_ACCENT_GREEN
        r.font.bold = True
        if i < 2:
            sep = specs_p.add_run(' │ ')
            sep.font.color.rgb = C_GRAY_400
            sep.font.size = Pt(9)

    div2 = cell.add_paragraph()
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div2.paragraph_format.space_before = Pt(10)
    div2.paragraph_format.space_after = Pt(4)
    div2_r = div2.add_run('━' * 50)
    div2_r.font.color.rgb = C_ACCENT_INDIGO
    div2_r.font.size = Pt(10)

    add_cover_para('6 Chapters · Complete Source Code · Visual Diagrams · Runbook', 9, False, C_GRAY_400, 4, 30)

    add_pagebreak(doc)

# ─── CHAPTER DIVIDER ─────────────────────────────────────────────────────────

def add_chapter_divider(doc, chapter_num, title, subtitle='', color=None):
    color = color or C_CHAPTER_HEADER
    tbl = doc.add_table(rows=1, cols=1)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(14)
    p0.paragraph_format.space_after = Pt(4)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f'CHAPTER {chapter_num}')
    r0.font.name = 'Calibri Light'
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFF)
    r0.font.bold = False

    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title.upper())
    r1.font.name = 'Calibri Light'
    r1.font.size = Pt(20)
    r1.font.color.rgb = C_WHITE
    r1.font.bold = True

    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(14)
        r2 = p2.add_run(subtitle)
        r2.font.name = 'Calibri Light'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFF)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── SECTION HEADING ─────────────────────────────────────────────────────────

def add_section_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = 'Calibri Light'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = C_ACCENT_INDIGO
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = C_ACCENT_BLUE
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = C_GRAY_700

    # Add underline bar for H1
    if level == 1:
        bar = doc.add_paragraph()
        bar.paragraph_format.space_before = Pt(0)
        bar.paragraph_format.space_after = Pt(8)
        bar_run = bar.add_run('▬' * 45)
        bar_run.font.color.rgb = C_ACCENT_INDIGO
        bar_run.font.size = Pt(7)

# ─── BODY TEXT ───────────────────────────────────────────────────────────────

def add_body(doc, text, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Pt(indent)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_GRAY_700

# ─── CALLOUT BOX ─────────────────────────────────────────────────────────────

def add_callout(doc, title, lines, accent=None, icon='▶'):
    accent = accent or C_ACCENT_BLUE
    tbl = doc.add_table(rows=1, cols=2)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)

    # Left accent stripe
    left = tbl.cell(0, 0)
    left.width = Inches(0.1)
    set_cell_bg(left, accent)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(4)

    # Right content
    right = tbl.cell(0, 1)
    right.width = Inches(6.4)
    set_cell_bg(right, RGBColor(0xF0, 0xF4, 0xFF))

    rp = right.paragraphs[0]
    rp.paragraph_format.space_before = Pt(6)
    rp.paragraph_format.space_after = Pt(2)
    rp.paragraph_format.left_indent = Pt(8)
    title_run = rp.add_run(f'{icon}  {title}')
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(10)
    title_run.font.bold = True
    title_run.font.color.rgb = accent

    for line in lines:
        lp2 = right.add_paragraph()
        lp2.paragraph_format.space_before = Pt(1)
        lp2.paragraph_format.space_after = Pt(1)
        lp2.paragraph_format.left_indent = Pt(8)
        lr = lp2.add_run(line)
        lr.font.name = 'Calibri'
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = C_GRAY_700

    # Bottom padding para
    bp = right.add_paragraph()
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── CODE BLOCK ──────────────────────────────────────────────────────────────

def add_code_block(doc, lines, title=''):
    if title:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(8)
        tp.paragraph_format.space_after = Pt(0)
        tr = tp.add_run(f'  {title}')
        tr.font.name = 'Calibri'
        tr.font.size = Pt(8.5)
        tr.font.bold = True
        tr.font.color.rgb = C_ACCENT_BLUE

    tbl = doc.add_table(rows=1, cols=1)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, C_CODE_BG)

    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(8)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        # Colorize keywords
        if line.strip().startswith('#'):
            run.font.color.rgb = RGBColor(0x6B, 0x73, 0x80)  # comment grey
        elif any(kw in line for kw in ['def ', 'class ', 'return ', 'import ', 'from ', 'async ', 'await ']):
            run.font.color.rgb = RGBColor(0xC0, 0x92, 0xFF)  # purple for keywords
        elif line.strip().startswith('//') or line.strip().startswith('const ') or line.strip().startswith('let '):
            run.font.color.rgb = RGBColor(0x7D, 0xD3, 0xFC)
        else:
            run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)

    # Padding after code
    ep = cell.add_paragraph()
    ep.paragraph_format.space_before = Pt(2)
    ep.paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── DATA TABLE ──────────────────────────────────────────────────────────────

def add_data_table(doc, headers, rows, accent=None):
    accent = accent or C_ACCENT_INDIGO
    col_count = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, accent)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = C_WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = C_GRAY_100 if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, C_GRAY_200, 4)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Pt(4) if ci == 0 else Pt(0)
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.color.rgb = C_GRAY_700

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ─── FLOW DIAGRAM (as a styled table grid) ───────────────────────────────────

def add_flow_node(doc, nodes, accent=None, direction='vertical'):
    """Render a list of labelled nodes as a vertical or horizontal flow."""
    accent = accent or C_ACCENT_BLUE
    if direction == 'vertical':
        for i, (label, sublabel) in enumerate(nodes):
            tbl = doc.add_table(rows=1, cols=1)
            table_no_borders(tbl)
            set_table_width(tbl, 5.0)
            cell = tbl.cell(0, 0)
            node_color = accent if i == 0 else (C_GRAY_800 if i % 2 == 0 else RGBColor(0x1E, 0x3A, 0x5F))
            set_cell_bg(cell, node_color)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(label)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = C_WHITE

            if sublabel:
                sp = cell.add_paragraph()
                sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sp.paragraph_format.space_before = Pt(1)
                sp.paragraph_format.space_after = Pt(5)
                sr = sp.add_run(sublabel)
                sr.font.name = 'Calibri'
                sr.font.size = Pt(8)
                sr.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFF)
            else:
                ep = cell.add_paragraph()
                ep.paragraph_format.space_before = Pt(0)
                ep.paragraph_format.space_after = Pt(5)

            if i < len(nodes) - 1:
                arrow_p = doc.add_paragraph()
                arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                arrow_p.paragraph_format.space_before = Pt(0)
                arrow_p.paragraph_format.space_after = Pt(0)
                ar = arrow_p.add_run('▼')
                ar.font.color.rgb = accent
                ar.font.size = Pt(14)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ─── HORIZONTAL PIPELINE FLOW ────────────────────────────────────────────────

def add_horizontal_flow(doc, stages, accent=None):
    """Render stages as a horizontal flow in one wide table."""
    accent = accent or C_ACCENT_BLUE
    n = len(stages)
    col_count = n * 2 - 1  # stages + arrows
    tbl = doc.add_table(rows=2, cols=col_count)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)

    for i, (label, color) in enumerate(stages):
        col = i * 2
        # Top row: stage box
        cell = tbl.cell(0, col)
        set_cell_bg(cell, color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(label)
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = C_WHITE

        # Arrow cell between stages
        if i < len(stages) - 1:
            arrow_cell = tbl.cell(0, col + 1)
            set_cell_bg(arrow_cell, C_GRAY_200)
            ap = arrow_cell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_before = Pt(6)
            ap.paragraph_format.space_after = Pt(6)
            ar = ap.add_run('▶')
            ar.font.color.rgb = accent
            ar.font.size = Pt(10)

    # Bottom row: phase number labels
    for i, (label, color) in enumerate(stages):
        col = i * 2
        cell = tbl.cell(1, col)
        set_cell_bg(cell, RGBColor(0xE0, 0xE7, 0xFF))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'PHASE {i+1}')
        run.font.name = 'Calibri'
        run.font.size = Pt(7)
        run.font.color.rgb = C_ACCENT_INDIGO

        if i < len(stages) - 1:
            ac = tbl.cell(1, col + 1)
            set_cell_bg(ac, C_GRAY_200)
            ac.paragraphs[0].paragraph_format.space_before = Pt(2)
            ac.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ─── HARDWARE TOPOLOGY DIAGRAM ────────────────────────────────────────────────

def add_hardware_topology(doc):
    add_section_heading(doc, '1.3 Workstation Hardware Topology', 2)
    add_body(doc, 'Data flows from NVMe storage through page-locked RAM into VRAM. The asyncio.Semaphore(2) token gate strictly limits concurrent GPU access to 2 sessions at all times.')

    # 3-row topology diagram
    tbl = doc.add_table(rows=7, cols=3)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)

    def fill_cell(row, col, text, bg, fg=C_WHITE, size=9, bold=False, align='center', sub=''):
        cell = tbl.cell(row, col)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = fg
        if sub:
            sp = cell.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(5)
            sr = sp.add_run(sub)
            sr.font.name = 'Consolas'
            sr.font.size = Pt(7.5)
            sr.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFF)
        else:
            ep = cell.add_paragraph()
            ep.paragraph_format.space_before = Pt(0)
            ep.paragraph_format.space_after = Pt(5)

    def fill_arrow(row, col, symbol='▼'):
        cell = tbl.cell(row, col)
        set_cell_bg(cell, C_WHITE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(symbol)
        run.font.color.rgb = C_ACCENT_BLUE
        run.font.size = Pt(14)

    # Row 0: NVMe header
    merge0 = tbl.cell(0, 0).merge(tbl.cell(0, 2))
    fill_cell(0, 0, '💾  NVMe SSD  —  SPOVNOB_intake/  &  pipeline_system_outputs/',
              RGBColor(0x1E, 0x3A, 0x5F), C_WHITE, 10, True)

    fill_arrow(1, 0, '▼'); fill_arrow(1, 1, '▼'); fill_arrow(1, 2, '▼')

    # Row 2: RAM sections
    fill_cell(2, 0, '🧠  512 GB ECC RAM', RGBColor(0x14, 0x53, 0x2D), C_WHITE, 9, True,
              sub='asyncio event loop\nWatchdog observer')
    fill_cell(2, 1, '🧵  12× CPU Workers', RGBColor(0x14, 0x53, 0x2D), C_WHITE, 9, True,
              sub='MediaPipe pool\nOpenFace MLT')
    fill_cell(2, 2, '📦  DataFrames', RGBColor(0x14, 0x53, 0x2D), C_WHITE, 9, True,
              sub='DynamicWindowEngine\nBaselineCalibrator')

    fill_arrow(3, 0, ' '); fill_arrow(3, 1, '▼'); fill_arrow(3, 2, ' ')

    # Row 4: PCIe label
    merge4 = tbl.cell(4, 0).merge(tbl.cell(4, 2))
    fill_cell(4, 0, '⚡  PCIe Gen 4 x16  —  pin_memory=True  —  non_blocking DMA Transfers',
              RGBColor(0x78, 0x35, 0x0F), C_WHITE, 9, True)

    fill_arrow(5, 0, ' '); fill_arrow(5, 1, '▼'); fill_arrow(5, 2, ' ')

    # Row 6: VRAM
    merge6 = tbl.cell(6, 0).merge(tbl.cell(6, 2))
    fill_cell(6, 0,
              '🎮  NVIDIA RTX 6000 Ada  —  48 GB VRAM  |  YOLOv8 TensorRT · HuBERT L7 · 2D CNN TFN\n'
              '             *** GUARDED BY: asyncio.Semaphore(2) — MAX 2 CONCURRENT GPU SESSIONS ***',
              RGBColor(0x1E, 0x1B, 0x4B), C_WHITE, 9, True)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ─── PIPELINE ORCHESTRATION FLOW ─────────────────────────────────────────────

def add_pipeline_flow(doc):
    add_body(doc, 'The MultimodalProductionOrchestrator executes 4 sequential phases per session:')
    stages = [
        ('PHASE 1\nVisual &\nAudio\nExtraction', RGBColor(0x1D, 0x4E, 0xD8)),
        ('PHASE 2\nRaw Feature\nCompilation\n30 FPS', RGBColor(0x0F, 0x76, 0x6E)),
        ('PHASE 3\nSliding Window\nAggregation\n2s / 1s stride', RGBColor(0x7C, 0x3A, 0xED)),
        ('PHASE 4\nBaseline\nCalibration\nZ-Score', RGBColor(0xB4, 0x5A, 0x09)),
    ]
    add_horizontal_flow(doc, stages, C_ACCENT_INDIGO)

# ─── HMM STATE MACHINE DIAGRAM ───────────────────────────────────────────────

def add_hmm_diagram(doc):
    add_body(doc, 'The HMM maintains two states with phase-modulated transition matrices. Higher-pressure phases relax the self-transition probability, making friction detection more sensitive:')

    tbl = doc.add_table(rows=3, cols=5)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)

    # State 0
    s0 = tbl.cell(1, 0)
    set_cell_bg(s0, C_ACCENT_GREEN)
    s0p = s0.paragraphs[0]
    s0p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s0p.paragraph_format.space_before = Pt(12)
    s0p.paragraph_format.space_after = Pt(2)
    s0r = s0p.add_run('STATE 0')
    s0r.font.name = 'Calibri'; s0r.font.bold = True; s0r.font.size = Pt(10); s0r.font.color.rgb = C_WHITE
    s0sub = s0.add_paragraph()
    s0sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s0sub.paragraph_format.space_before = Pt(0)
    s0sub.paragraph_format.space_after = Pt(12)
    s0subr = s0sub.add_run('Stable Context')
    s0subr.font.name = 'Calibri'; s0subr.font.size = Pt(8); s0subr.font.color.rgb = C_WHITE

    # Arrows
    for r in [0, 1, 2]:
        for c in [1, 3]:
            ac = tbl.cell(r, c)
            set_cell_bg(ac, C_WHITE)
    mid_top = tbl.cell(0, 2)
    set_cell_bg(mid_top, C_WHITE)
    mtp = mid_top.paragraphs[0]
    mtp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mtp.paragraph_format.space_before = Pt(4)
    mtr = mtp.add_run('P(0→1)\nfriction onset')
    mtr.font.name = 'Calibri'; mtr.font.size = Pt(7.5); mtr.font.color.rgb = C_ACCENT_AMBER

    mid_bot = tbl.cell(2, 2)
    set_cell_bg(mid_bot, C_WHITE)
    mbp = mid_bot.paragraphs[0]
    mbp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mbp.paragraph_format.space_before = Pt(4)
    mbr = mbp.add_run('P(1→0)\nrecovery')
    mbr.font.name = 'Calibri'; mbr.font.size = Pt(7.5); mbr.font.color.rgb = C_ACCENT_GREEN

    # State 1
    s1 = tbl.cell(1, 4)
    set_cell_bg(s1, C_ACCENT_RED)
    s1p = s1.paragraphs[0]
    s1p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1p.paragraph_format.space_before = Pt(12)
    s1p.paragraph_format.space_after = Pt(2)
    s1r = s1p.add_run('STATE 1')
    s1r.font.name = 'Calibri'; s1r.font.bold = True; s1r.font.size = Pt(10); s1r.font.color.rgb = C_WHITE
    s1sub = s1.add_paragraph()
    s1sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1sub.paragraph_format.space_before = Pt(0)
    s1sub.paragraph_format.space_after = Pt(12)
    s1subr = s1sub.add_run('High Cognitive Load')
    s1subr.font.name = 'Calibri'; s1subr.font.size = Pt(8); s1subr.font.color.rgb = C_WHITE

    # Middle column: arrows
    arrowmid = tbl.cell(1, 2)
    set_cell_bg(arrowmid, C_WHITE)
    amp = arrowmid.paragraphs[0]
    amp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    amp.paragraph_format.space_before = Pt(10)
    amr1 = amp.add_run('━━━▶\n◀━━━')
    amr1.font.name = 'Consolas'; amr1.font.size = Pt(10); amr1.font.color.rgb = C_ACCENT_PURPLE

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ─── TFN TENSOR MAP ──────────────────────────────────────────────────────────

def add_tfn_tensor_map(doc):
    tbl = doc.add_table(rows=5, cols=3)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)

    shapes = [
        ('Visual Input\n[B, 114]', C_ACCENT_BLUE, 'Linear(114→70)\nXavier Init'),
        ('Acoustic Input\n[B, 20]', C_ACCENT_GREEN, 'Linear(20→70)\nXavier Init'),
    ]
    labels = ['V_proj [B,70] + bias → [B,71]', 'A_proj [B,70] + bias → [B,71]']
    fused = 'torch.bmm([B,71,1] × [B,1,71]) → interaction [B,71,71]\n→ unsqueeze(1) → [B, 1, 71, 71]  ←  CNN Input'

    for i, ((label, color, sublabel), proj_label) in enumerate(zip(shapes, labels)):
        row = i * 2
        # Input box
        inp = tbl.cell(row, 0)
        set_cell_bg(inp, color)
        ip = inp.paragraphs[0]
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.paragraph_format.space_before = Pt(8)
        ip.paragraph_format.space_after = Pt(2)
        ir = ip.add_run(label)
        ir.font.name = 'Consolas'; ir.font.size = Pt(9); ir.font.bold = True; ir.font.color.rgb = C_WHITE

        # Arrow
        ar_cell = tbl.cell(row, 1)
        set_cell_bg(ar_cell, C_WHITE)
        arp = ar_cell.paragraphs[0]
        arp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        arp.paragraph_format.space_before = Pt(8)
        arr_r = arp.add_run(f'─────▶\n{sublabel}')
        arr_r.font.name = 'Calibri'; arr_r.font.size = Pt(8); arr_r.font.color.rgb = C_GRAY_400

        # Projected
        proj = tbl.cell(row, 2)
        set_cell_bg(proj, RGBColor(0x1E, 0x3A, 0x5F))
        pp = proj.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(8)
        pp.paragraph_format.space_after = Pt(2)
        pr = pp.add_run(proj_label)
        pr.font.name = 'Consolas'; pr.font.size = Pt(8); pr.font.color.rgb = RGBColor(0x7D, 0xD3, 0xFC)

        if row + 1 < 5:
            for c in range(3):
                mid = tbl.cell(row + 1, c)
                set_cell_bg(mid, C_WHITE)
                mp = mid.paragraphs[0]
                mp.paragraph_format.space_before = Pt(2)
                mp.paragraph_format.space_after = Pt(2)

    # Final fused row
    fuse_merge = tbl.cell(4, 0).merge(tbl.cell(4, 2))
    set_cell_bg(fuse_merge, RGBColor(0x1E, 0x1B, 0x4B))
    fp = fuse_merge.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(8)
    fp.paragraph_format.space_after = Pt(8)
    fr = fp.add_run(fused)
    fr.font.name = 'Consolas'; fr.font.size = Pt(8.5); fr.font.bold = True
    fr.font.color.rgb = RGBColor(0xA5, 0xB4, 0xFC)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ─── CNN BACKBONE DIAGRAM ────────────────────────────────────────────────────

def add_cnn_diagram(doc):
    layers = [
        ('INPUT  [B, 1, 71, 71]',         RGBColor(0x1E, 0x3A, 0x5F), 'TFN interaction image — single channel'),
        ('Conv2d(1→32, k=3, p=1)',         RGBColor(0x1D, 0x4E, 0xD8), 'BatchNorm2d(32) → ReLU → MaxPool2d(2)'),
        ('[B, 32, 35, 35]',               RGBColor(0x07, 0x3A, 0x63), '32 feature maps at 35×35 resolution'),
        ('Conv2d(32→64, k=3, p=1)',        RGBColor(0x1D, 0x4E, 0xD8), 'BatchNorm2d(64) → ReLU → MaxPool2d(2)'),
        ('[B, 64, 17, 17]',               RGBColor(0x07, 0x3A, 0x63), '64 feature maps at 17×17 resolution'),
        ('Conv2d(64→128, k=3, p=1)',       RGBColor(0x1D, 0x4E, 0xD8), 'BatchNorm2d(128) → ReLU → AdaptiveAvgPool(4)'),
        ('[B, 128, 4, 4]',                RGBColor(0x07, 0x3A, 0x63), '128 feature maps at 4×4 resolution'),
        ('Flatten  →  [B, 2048]',         RGBColor(0x4A, 0x1D, 0x96), 'Flatten all spatial dimensions'),
        ('Linear(2048→256) → ReLU → Dropout(0.4)', RGBColor(0x4A, 0x1D, 0x96), 'FC hidden layer with dropout'),
        ('Linear(256→2)  →  LOGITS [B, 2]', RGBColor(0x0F, 0x76, 0x6E), 'State 0: Stable  |  State 1: Friction'),
    ]
    add_flow_node(doc, [(l, s) for l, _, s in layers], accent=C_ACCENT_BLUE)

# ─── VITERBI TRELLIS DIAGRAM ─────────────────────────────────────────────────

def add_viterbi_diagram(doc):
    add_body(doc, 'The Viterbi algorithm computes the globally optimal state path across T timesteps using dynamic programming on a log-space trellis:')
    tbl = doc.add_table(rows=5, cols=6)
    table_no_borders(tbl)
    set_table_width(tbl, 6.5)

    headers = ['t=0 (Init)', 't=1', 't=2', '…', 't=T-2', 't=T-1 (Term)']
    for ci, h in enumerate(headers):
        cell = tbl.cell(0, ci)
        set_cell_bg(cell, C_GRAY_800)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.font.name = 'Consolas'; run.font.size = Pt(8); run.font.color.rgb = C_GRAY_200

    # State rows
    for ri, (state, color, label) in enumerate([
        ('V[t, 0]\nSTABLE', C_ACCENT_GREEN, 'log π(0) + log B[0,0]'),
        ('━━━━━━', C_WHITE, ''),
        ('V[t, 1]\nFRICTION', C_ACCENT_RED, 'log π(1) + log B[0,1]'),
    ]):
        for ci in range(6):
            cell = tbl.cell(ri + 1, ci)
            if state == '━━━━━━':
                set_cell_bg(cell, C_WHITE)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
            else:
                set_cell_bg(cell, color if ci == 0 else RGBColor(0x1E, 0x29, 0x3B))
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                text = state if ci == 0 else (label if ci == 1 else 'max(V[t-1,:] + log_A[:,s])\n+ log B[t,s]')
                run = p.add_run(text)
                run.font.name = 'Consolas'; run.font.size = Pt(7.5)
                run.font.color.rgb = C_WHITE if ci == 0 else RGBColor(0x7D, 0xD3, 0xFC)

    # Backtrack row
    bt_merge = tbl.cell(4, 0).merge(tbl.cell(4, 5))
    set_cell_bg(bt_merge, RGBColor(0x1E, 0x1B, 0x4B))
    bp = bt_merge.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(5)
    bp.paragraph_format.space_after = Pt(5)
    br_run = bp.add_run('BACKTRACKING: optimal_path[t] = backptr[t+1, optimal_path[t+1]]  →  flicker-free session_report.json')
    br_run.font.name = 'Consolas'; br_run.font.size = Pt(8); br_run.font.color.rgb = RGBColor(0xA5, 0xB4, 0xFC)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ─── ZUSTAND DATA FLOW ───────────────────────────────────────────────────────

def add_zustand_flow(doc):
    nodes = [
        ('VIDEO ELEMENT  onTimeUpdate()', 'HTMLVideoElement native event — fires at native frame rate'),
        ('useStore.getState().setGlobalTimeMs(ms)', 'Direct Zustand state mutation — NO React render triggered'),
        ('useStore.subscribe(state => ...)', '60Hz transient callback — bypasses virtual DOM reconciler'),
        ('echartsInstance.dispatchAction({ type: "showTip" })', 'Direct Canvas API call — crosshair moved in raw WebGL layer'),
        ('echartsInstance.setOption({ series: [{ value: auValues }] })', 'RadarChart AU polygon updated — zero React state change'),
    ]
    add_flow_node(doc, nodes, accent=C_ACCENT_INDIGO)

# ─── MAIN BUILD ──────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ── COVER ──────────────────────────────────────────────────────────────
    build_cover(doc)

    # ── TABLE OF CONTENTS ──────────────────────────────────────────────────
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_before = Pt(10)
    toc_title.paragraph_format.space_after = Pt(10)
    t_run = toc_title.add_run('TABLE OF CONTENTS')
    t_run.font.name = 'Calibri Light'; t_run.font.size = Pt(18); t_run.font.bold = True
    t_run.font.color.rgb = C_ACCENT_INDIGO

    toc_items = [
        ('Chapter 1', 'Intellectual Defense Posture, Timelines & Workstation Ceilings'),
        ('Chapter 2', 'The Automated Intake Refinery & System-Level Resource Locks'),
        ('Chapter 3', 'Dual-Engine Spatial Extraction & Audio Isolation Biometrics'),
        ('Chapter 4', 'Sliding Window Statistics & Automated Forensic Label Matching'),
        ('Chapter 5', 'The Core Predictive Ensemble — TFN · CNN · HMM Viterbi'),
        ('Chapter 6', 'Sidecar Broadcasting, Frontend Compiling & Subscription Routing'),
        ('Appendix', 'Runbook & Operational Orchestration Playbook'),
    ]
    colors = [C_ACCENT_BLUE, C_ACCENT_GREEN, C_ACCENT_PURPLE, C_ACCENT_AMBER,
              C_ACCENT_INDIGO, C_ACCENT_RED, C_GRAY_700]

    for i, ((num, title), color) in enumerate(zip(toc_items, colors)):
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(3)
        tp.paragraph_format.space_after = Pt(3)
        num_run = tp.add_run(f'  {num}  ')
        num_run.font.name = 'Calibri'; num_run.font.size = Pt(11); num_run.font.bold = True
        num_run.font.color.rgb = color
        title_run = tp.add_run(title)
        title_run.font.name = 'Calibri'; title_run.font.size = Pt(11)
        title_run.font.color.rgb = C_GRAY_700

    add_pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER 1
    # ══════════════════════════════════════════════════════════════════════
    add_chapter_divider(doc, 1,
        'Intellectual Defense Posture, Timelines & Workstation Ceilings',
        'Scientific core, VFR rejection, NaN philosophy, hardware topology',
        RGBColor(0x1E, 0x40, 0xAF))

    add_section_heading(doc, '1.1 Scientific Core vs. Commercial Bias', 2)
    add_body(doc,
        'SPOVNOB is engineered as a Temporal Behavioral Anomaly Detection and Cognitive Load Analysis '
        'platform. It produces a continuous, normalized behavioral deviation signal keyed to absolute '
        'millisecond time coordinates — not frame indices.')

    add_callout(doc, 'Why Absolute Millisecond Timelines?', [
        'Variable Frame Rate (VFR) video has per-frame intervals ranging from 28ms to 40ms.',
        'Frame counter 900 maps to 30.0s on a CFR clock but 34.7s on the actual media clock.',
        'Dropped frames create invisible 100ms behavioral voids under frame-index tracking.',
        'cv2.CAP_PROP_POS_MSEC reads the container\'s actual PTS — the single authoritative clock.',
        'All downstream merges, windows, and acoustic alignment are keyed to this clock.',
    ], C_ACCENT_BLUE, '⚡')

    add_data_table(doc,
        ['Failure Mode', 'Frame-Index Result', 'Millisecond Timeline Fix'],
        [
            ['VFR timestamp drift', 'Frame 900 ≠ 30.0 s', 'PTS-anchored, exact'],
            ['Dropped frames', 'Invisible 100ms voids', 'NaN preserved, visible'],
            ['Audio/video desync', 'Gross misalignment', '≤ ±30ms tolerance'],
            ['HuBERT window slip', 'Wrong audio slice', 'Exact ms window boundary'],
        ], C_ACCENT_BLUE)

    add_section_heading(doc, '1.2 The Philosophy of the Natural NaN State', 2)
    add_body(doc,
        'When face-lock fails (subject covers face, extreme rotation, tracker loss), the pipeline '
        'records np.nan — NOT zero, NOT interpolated. This is a critical design choice:')

    add_callout(doc, 'NaN vs. Zero — Engineering Rationale', [
        'Zero = a measured, confirmed absence of signal (e.g., eye fully open, EAR=0.0).',
        'NaN  = no measurement at all — occlusion itself is a behavioral data point.',
        'Filling NaN→0 would create ghost signals: "spike to zero" = false deviation anomaly.',
        'Downstream: ConfidenceWeightedCrossEntropy downweights high-NaN windows in gradient updates.',
        'NaN→0.0 clamping only happens at the Z-score tensor extraction gate — mathematically safe.',
        '"0.0 in Z-score space" = exactly at baseline mean — most conservative possible imputation.',
    ], C_ACCENT_AMBER, '⚠')

    add_code_block(doc, [
        '# analytics/predictive_engine.py — lines 248–249',
        '# Clamp happens ONLY after Z-score normalization, at the final tensor extraction gate',
        'visual_clean  = np.nan_to_num(visual_raw,   nan=0.0, posinf=0.0, neginf=0.0)',
        'acoustic_clean = np.nan_to_num(acoustic_raw, nan=0.0, posinf=0.0, neginf=0.0)',
    ], 'Confidence Imputation Gate — predictive_engine.py:248')

    add_hardware_topology(doc)
    add_pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER 2
    # ══════════════════════════════════════════════════════════════════════
    add_chapter_divider(doc, 2,
        'The Automated Intake Refinery & System-Level Resource Locks',
        'Watchdog · fcntl · Ledger · SIGINT/SIGTERM · Spawn VRAM isolation',
        RGBColor(0x06, 0x57, 0x2A))

    add_section_heading(doc, '2.1 File-System Watchdog & fcntl Atomicity', 2)
    add_body(doc, 'The batch daemon monitors SPOVNOB_intake/ and gates ingestion on file lock availability:')

    # Watchdog sequence flow
    seq_nodes = [
        ('FS: FileCreated(session_profile.json)', 'SPOVNOB_intake/<SESSION_ID>/ detected by watchdog'),
        ('IntakeManifestHandler.on_created()', 'asyncio.run_coroutine_threadsafe → queue.put(path)'),
        ('asyncio.Queue dequeued by main_loop()', 'Profile JSON parsed — session_id extracted'),
        ('fcntl.flock(LOCK_EX | LOCK_NB) on .mp4 + .wav', 'Non-blocking exclusive lock check — defer if IOError'),
        ('asyncio.Semaphore(2) acquired', 'Blocks if 2 GPU sessions already active'),
        ('spawn_ctx.Process.start() → GPU Worker', 'Fresh CUDA context in isolated subprocess'),
        ('process.join() via run_in_executor', 'Event loop free during GPU processing'),
        ('ELAN label injection → set_state("COMPLETED")', 'Atomic ledger flush via os.replace()'),
    ]
    add_flow_node(doc, seq_nodes, C_ACCENT_GREEN)

    add_code_block(doc, [
        '# app/batch_daemon.py — lines 264–276',
        'def is_file_transfer_complete(file_path: Path) -> bool:',
        '    """',
        '    Attempts a non-blocking exclusive lock on the target file.',
        '    If lock succeeds: file is fully written, not held by rsync/scp/NFS.',
        '    LOCK_NB → returns EWOULDBLOCK immediately instead of blocking.',
        '    """',
        '    try:',
        '        with open(file_path, "rb") as f:',
        '            fcntl.flock(f.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)',
        '            fcntl.flock(f.fileno(), fcntl.LOCK_UN)',
        '        return True',
        '    except (IOError, OSError):',
        '        return False',
    ], 'fcntl Non-Blocking Lock — batch_daemon.py:264')

    add_section_heading(doc, '2.2 File Descriptor Exhaustion Shield', 2)
    add_code_block(doc, [
        '# app/batch_daemon.py — lines 39–52 (runs at module import time)',
        'def maximize_file_descriptors():',
        '    soft, hard = resource.getrlimit(resource.RLIMIT_NOFILE)',
        '    if soft < hard:',
        '        resource.setrlimit(resource.RLIMIT_NOFILE, (hard, hard))',
        '        logging.info(f"RLIMIT_NOFILE raised: {soft} → {hard}")',
        '',
        'maximize_file_descriptors()  # Runs before ANY other code',
    ], 'FD Exhaustion Shield — batch_daemon.py:39')

    add_section_heading(doc, '2.3 Ledger State Machine & Atomic Write', 2)

    # State machine flow
    sm_nodes = [
        ('QUEUED', 'Session manifest detected — files verified by fcntl'),
        ('TENSORRT_ACTIVE', 'GPU semaphore acquired — spawn subprocess running'),
        ('MATH_NORMALIZATION', 'Subprocess exit code 0 — ELAN label injection'),
        ('COMPLETED  /  FAILED  /  INTERRUPTED', 'Terminal state — atomic ledger flush via os.replace()'),
    ]
    add_flow_node(doc, sm_nodes, C_ACCENT_AMBER)

    add_code_block(doc, [
        '# app/batch_daemon.py — lines 106–119',
        'def flush(self):',
        '    """Atomic write: write to temp file, then os.replace to prevent corruption."""',
        '    tmp_fd, tmp_path = tempfile.mkstemp(',
        '        dir=str(self.ledger_path.parent), suffix=".tmp"',
        '    )  # Same directory → same mount point → atomic rename',
        '    try:',
        '        with os.fdopen(tmp_fd, "w") as f:',
        '            json.dump(self.entries, f, indent=2)',
        '        os.replace(tmp_path, str(self.ledger_path))  # ATOMIC on POSIX',
        '    except Exception:',
        '        if os.path.exists(tmp_path):',
        '            os.unlink(tmp_path)',
        '        raise',
    ], 'Atomic Ledger Flush — batch_daemon.py:106')

    add_code_block(doc, [
        '# app/batch_daemon.py — lines 500–517',
        'def install(self):',
        '    signal.signal(signal.SIGINT,  self._handle)',
        '    signal.signal(signal.SIGTERM, self._handle)',
        '',
        'def _handle(self, signum, frame):',
        '    if self._triggered: return',
        '    self._triggered = True',
        '    sig_name = signal.Signals(signum).name',
        '    logger.warning(f"🛑 Received {sig_name}. Initiating graceful shutdown...")',
        '    self.orchestrator.request_shutdown()',
        '    self.orchestrator.terminate_active()  # SIGTERM → SIGKILL on subprocesses',
        '    self.ledger.mark_interrupted()         # Stamps all active → INTERRUPTED',
        '    sys.exit(0)',
    ], 'SIGINT / SIGTERM Signal Trapping — batch_daemon.py:500')

    add_pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER 3
    # ══════════════════════════════════════════════════════════════════════
    add_chapter_divider(doc, 3,
        'Dual-Engine Spatial Extraction & Audio Isolation Biometrics',
        'YOLOv8 · FaceLock · MediaPipe · OpenFace AUs · HuBERT · Lip-Sync Gate',
        RGBColor(0x4A, 0x1D, 0x96))

    add_section_heading(doc, '3.1 Spatial Visual Extraction — OpenFace Action Units', 2)
    add_data_table(doc,
        ['AU Column', 'OpenFace Source', 'Semantic Meaning', 'Failure Value'],
        [
            ['AU1', 'AU01_r intensity', 'Inner Brow Raiser — fear/surprise', 'np.nan'],
            ['AU2', 'AU02_r intensity', 'Outer Brow Raiser — surprise/disbelief', 'np.nan'],
            ['AU4', 'AU04_r intensity', 'Brow Lowerer — confusion/anger', 'np.nan'],
            ['AU6', 'AU06_r intensity', 'Cheek Raiser — genuine smile marker', 'np.nan'],
            ['AU9', 'AU09_r intensity', 'Nose Wrinkler — disgust leakage', 'np.nan'],
            ['AU12', 'AU12_r intensity', 'Lip Corner Puller — Duchenne smile', 'np.nan'],
            ['AU25', 'AU25_r intensity', 'Lips Part — speech onset', 'np.nan'],
            ['AU26', 'AU26_r intensity', 'Jaw Drop — surprise/effort', 'np.nan'],
        ], C_ACCENT_PURPLE)

    add_section_heading(doc, '3.2 3D Kinematic Feature Engineering', 2)
    add_code_block(doc, [
        '# main_pipeline.py — lines 76–99',
        '# Full 3D Euclidean distance — includes depth component (proximity bias cancelled)',
        'df_pose["left_hand_face_distance"] = np.sqrt(',
        '    (df_pose["left_wrist_x"] - df_pose["nose_x"])**2 +',
        '    (df_pose["left_wrist_y"] - df_pose["nose_y"])**2 +',
        '    (df_pose["left_wrist_z"] - df_pose["nose_z"])**2',
        ')',
        '',
        '# Wrist velocity — first derivative, NaN propagates on tracking gaps',
        'df_pose["left_wrist_velocity"] = np.sqrt(',
        '    df_pose["left_wrist_x"].diff()**2 +',
        '    df_pose["left_wrist_y"].diff()**2 +',
        '    df_pose["left_wrist_z"].diff()**2',
        ')',
        '',
        '# AU onset velocity — detects micro-expression speed (genuine: 250-500ms)',
        'au_columns = ["AU1", "AU2", "AU4", "AU6", "AU9", "AU12", "AU25", "AU26"]',
        'for au in au_columns:',
        '    fused_frames[f"{au}_velocity"] = fused_frames[au].diff().fillna(0)',
    ], 'Vectorized Kinematics — main_pipeline.py:76')

    add_section_heading(doc, '3.3 Lip-Sync Cocktail Party Gatekeeper', 2)
    add_data_table(doc,
        ['Lip State', 'Audio State', 'mismatch_incongruence', 'silent_speech', 'Interpretation'],
        [
            ['Closed', 'Silent', '0.0', '0.0', 'Target listening — normal'],
            ['Open', 'Active', '0.0', '0.0', 'Target speaking — normal'],
            ['Closed', 'Active', '1.0', '0.0', 'Background speaker contamination'],
            ['Open', 'Silent', '0.0', '1.0', 'Subvocalization / micro-articulation'],
        ], C_ACCENT_RED)

    add_section_heading(doc, '3.4 Joint Confidence Vector w_t', 2)
    add_callout(doc, 'Confidence Product: w_t = c_yolo × c_facelock × c_landmark × c_diarizer', [
        'If ANY single tracker fails (value = 0.0) → product collapses to 0.0.',
        'Zero-weight windows: zero gradient contribution in ConfidenceWeightedCrossEntropy.',
        'diarizer_conf = 0.90 if speaking, 1.00 if silent — speaking windows carry slight uncertainty.',
        'All NaN conf values filled with 0.0 before product to ensure strict safety.',
    ], C_ACCENT_GREEN, '🔒')

    add_pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER 4
    # ══════════════════════════════════════════════════════════════════════
    add_chapter_divider(doc, 4,
        'Sliding Window Statistics & Automated Forensic Label Matching',
        'Rolling descriptors · Z-score regularization · ELAN binary search',
        RGBColor(0x92, 0x40, 0x09))

    add_section_heading(doc, '4.1 The Rolling Descriptor Engine', 2)
    add_data_table(doc,
        ['Descriptor', 'Equation', 'Column Suffix', 'Semantic Use'],
        [
            ['Mean', '(1/n) Σ vᵢ', '_mean', 'Sustained behavioral level'],
            ['Peak / Max', 'max(v₀…vₙ)', '_max', 'Worst-case extrema within window'],
            ['Variance', '(1/n) Σ (vᵢ-μ)²', '_var', 'Signal instability / tremor'],
            ['Velocity', '(1/n) Σ |vᵢ - vᵢ₋₁|', '_velocity_mean', 'Rate of change / onset speed'],
            ['Count', 'Σ (~isnan(vᵢ))', '_count / fill_rate', 'Data density (tracking quality)'],
        ], C_ACCENT_AMBER)

    add_section_heading(doc, '4.2 Within-Subject Baseline Z-Score Regularization', 2)
    add_body(doc, 'BaselineCalibrator isolates the first 30,000ms of session data as the unpolluted neutral phase, then normalizes all windows against it:')

    add_code_block(doc, [
        '# analytics/baseline_calibrator.py — lines 96–131',
        '# Step 3: Compute baseline statistics',
        'baseline_mean = baseline_df[feature_cols].mean()',
        'baseline_std  = baseline_df[feature_cols].std()',
        '',
        '# Guard: zero std → NaN (NOT epsilon) — a constant feature during baseline',
        '# should never produce finite z-scores; infinite deviation is represented as NaN.',
        'baseline_std = baseline_std.replace(0, np.nan)',
        '',
        '# Step 4: Z-score normalize ALL windows (including baseline itself)',
        'df_calibrated[feature_cols] = (df[feature_cols] - baseline_mean) / baseline_std',
        '',
        '# Step 5: Per-window deviation magnitude (L2 norm of all z-scores)',
        '# High deviation_magnitude → subject behaving unusually vs. their personal baseline',
        'z_scores = df_calibrated[feature_cols]',
        "df_calibrated['deviation_magnitude'] = np.sqrt((z_scores ** 2).sum(axis=1))",
        '',
        '# Step 6: Percentile rank (0–1 scale) for pre-computed anomaly indicator',
        "df_calibrated['deviation_percentile'] = (",
        "    df_calibrated['deviation_magnitude'].rank(pct=True, na_option='keep')",
        ')',
    ], 'BaselineCalibrator.calibrate() — baseline_calibrator.py:96')

    add_section_heading(doc, '4.3 ELAN Annotation Binary Interval Search', 2)
    add_body(doc, 'ContextMapper uses bisect.bisect_right() for O(log N) interval lookup — ≤8 comparisons per window regardless of session length:')

    add_code_block(doc, [
        '# analytics/context_mapper.py — lines 49–69',
        'def lookup(self, timestamp_ms: float):',
        '    """O(log N) interval search. Returns (phase_label, question_id, elapsed_ms)"""',
        '    if not self.starts:',
        '        return np.nan, -1, np.nan',
        '    ',
        '    # bisect_right: insertion point AFTER all entries <= timestamp_ms',
        '    # idx - 1 = latest interval that started before or exactly at timestamp_ms',
        '    idx = bisect.bisect_right(self.starts, timestamp_ms) - 1',
        '    ',
        '    if idx >= 0:',
        '        start_ms, end_ms, phase_label, question_id = self.intervals[idx]',
        '        if timestamp_ms < end_ms:  # Confirm timestamp is INSIDE the interval',
        '            phase_elapsed_ms = float(timestamp_ms - start_ms)',
        '            return phase_label, question_id, phase_elapsed_ms',
        '    ',
        '    return np.nan, -1, np.nan',
    ], 'ContextMapper.lookup() — O(log N) binary search — context_mapper.py:49')

    add_pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER 5
    # ══════════════════════════════════════════════════════════════════════
    add_chapter_divider(doc, 5,
        'The Core Predictive Ensemble',
        'Tensor Fusion Network · 2D CNN Classifier · Context-Modulated HMM · Viterbi Decoder',
        RGBColor(0x1E, 0x1B, 0x4B))

    add_section_heading(doc, '5.1 CMU-Style Tensor Fusion Block — V ⊗ A', 2)
    add_body(doc, 'The TFN projects visual (114-dim) and acoustic (20-dim) vectors into a shared 70-dim latent space, appends a bias scalar, then computes the Cartesian Outer Product to produce a unified 71×71 interaction image:')
    add_tfn_tensor_map(doc)

    add_code_block(doc, [
        '# analytics/predictive_engine.py — lines 306–335',
        'def forward(self, visual, acoustic):',
        '    batch_size = visual.shape[0]',
        '    v_proj = self.visual_proj(visual)       # [B, 70]',
        '    a_proj = self.acoustic_proj(acoustic)   # [B, 70]',
        '',
        '    # Append bias scalar 1.0 — preserves unimodal features alongside cross-modal',
        '    ones = torch.ones(batch_size, 1, device=visual.device, dtype=visual.dtype)',
        '    v_augmented = torch.cat([v_proj, ones], dim=1)   # [B, 71]',
        '    a_augmented = torch.cat([a_proj, ones], dim=1)   # [B, 71]',
        '',
        '    # Cartesian Outer Product V ⊗ A via batch matrix multiply',
        '    v_col = v_augmented.unsqueeze(2)    # [B, 71,  1]',
        '    a_row = a_augmented.unsqueeze(1)    # [B,  1, 71]',
        '    interaction = torch.bmm(v_col, a_row)   # [B, 71, 71] — 5,041 nodes',
        '    return interaction.unsqueeze(1)         # [B,  1, 71, 71] ← CNN input',
    ], 'TensorFusionBlock.forward() — predictive_engine.py:306')

    add_section_heading(doc, '5.2 2D Convolutional Interaction Classifier Architecture', 2)
    add_cnn_diagram(doc)

    add_code_block(doc, [
        '# analytics/predictive_engine.py — lines 475–499',
        'class ConfidenceWeightedCrossEntropy(nn.Module):',
        '    """',
        '    Formula: L = -Σ(wᵢ · CE(logitsᵢ, targetsᵢ)) / Σ(wᵢ)',
        '    Where wᵢ = normalized cumulative_confidence for window i.',
        '    Windows with low tracker confidence → proportionally smaller gradient.',
        '    """',
        '    def __init__(self):',
        '        super().__init__()',
        '        self.base_loss = nn.CrossEntropyLoss(reduction="none")',
        '',
        '    def forward(self, logits, targets, confidence_weights):',
        '        per_sample_loss = self.base_loss(logits, targets)  # [B] — unreduced',
        '        weighted_loss   = per_sample_loss * confidence_weights  # [B] — weighted',
        '        weight_sum      = confidence_weights.sum().clamp(min=1e-9)',
        '        return weighted_loss.sum() / weight_sum  # Normalized weighted mean',
    ], 'ConfidenceWeightedCrossEntropy — predictive_engine.py:475')

    add_section_heading(doc, '5.3 Context-Modulated HMM — Phase Transition Profiles', 2)
    add_hmm_diagram(doc)

    add_data_table(doc,
        ['Phase', 'P(Stable→Stable)', 'P(Stable→Friction)', 'P(Friction→Friction)', 'P(Friction→Stable)'],
        [
            ['baseline_neutral', '0.98', '0.02', '0.85', '0.15'],
            ['briefing_instruction', '0.97', '0.03', '0.85', '0.15'],
            ['question_delivery', '0.90', '0.10', '0.88', '0.12'],
            ['subject_response', '0.80', '0.20', '0.92', '0.08'],
            ['investigative_confrontation', '0.75 ⬇ SENSITIVE', '0.25', '0.94 ⬆ STICKY', '0.06'],
            ['(default / unknown)', '0.85', '0.15', '0.90', '0.10'],
        ], C_ACCENT_INDIGO)

    add_section_heading(doc, '5.4 Viterbi Log-Space Trellis & Backtracking', 2)
    add_viterbi_diagram(doc)

    add_code_block(doc, [
        '# analytics/predictive_engine.py — lines 762–784',
        '# Initialization t=0',
        'for s in range(NUM_STATES):',
        '    V[0, s] = np.log(self.pi[s] + 1e-300) + np.log(emissions[0, s])',
        '',
        '# Recursion t=1..T-1',
        'for t in range(1, T):',
        '    phase = context_phases[t]',
        '    A     = self._get_transition_matrix(phase)  # Dynamic per-phase matrix',
        '    log_A = np.log(A + 1e-300)',
        '    for s in range(NUM_STATES):',
        '        candidates = V[t - 1, :] + log_A[:, s]',
        '        best_prev  = np.argmax(candidates)',
        '        V[t, s]    = candidates[best_prev] + np.log(emissions[t, s])',
        '        backptr[t, s] = best_prev',
        '',
        '# Backtracking — global optimal path reconstruction',
        'optimal_path[T - 1] = np.argmax(V[T - 1, :])',
        'for t in range(T - 2, -1, -1):',
        '    optimal_path[t] = backptr[t + 1, optimal_path[t + 1]]',
    ], 'Viterbi Decoding Loop — predictive_engine.py:762')

    add_pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER 6
    # ══════════════════════════════════════════════════════════════════════
    add_chapter_divider(doc, 6,
        'Sidecar Broadcasting, Frontend & Subscription Routing',
        'FastAPI · SSE · SPA Catch-All · Zustand Transient Store · ECharts LTTB',
        RGBColor(0x0F, 0x5E, 0x7E))

    add_section_heading(doc, '6.1 FastAPI Monolithic Sidecar — Endpoint Map', 2)
    add_data_table(doc,
        ['Method', 'Endpoint', 'Returns', 'Key Design Decision'],
        [
            ['GET', '/api/sessions', 'JSON list of metadata.json', 'secure_resolve() path traversal guard'],
            ['GET', '/api/data/{session_id}', 'orient="split" JSON', '60-70% smaller than orient="records"'],
            ['GET', '/api/video/{session_id}', 'HTTP 206 byte-range MP4', '4MB chunks, Accept-Ranges streaming'],
            ['GET', '/api/factory/status', 'Ledger snapshot JSON', 'Point-in-time ledger read'],
            ['GET', '/api/factory/stream', 'SSE text/event-stream', '2s poll, delta-only pushes'],
            ['GET', '/assets/{file_path}', 'Static JS/CSS bundle file', 'Traversal-safe FRONTEND_DIST guard'],
            ['GET', '/{full_path:path}', 'index.html (SPA fallback)', 'Eliminates F5 refresh 404 crashes'],
        ], C_ACCENT_BLUE)

    add_code_block(doc, [
        '# app/server.py — lines 227–233',
        '@app.get("/{full_path:path}")',
        'def catch_all(full_path: str):',
        '    """Fallback handler to prevent SPA Sub-Path Refresh Trap"""',
        '    index_file = (FRONTEND_DIST / "index.html").resolve()',
        '    if not index_file.exists():',
        '        return JSONResponse({"error": "Run npm run build in frontend/"}, status_code=404)',
        '    return FileResponse(index_file)',
        '',
        '# WHY: React Router uses window.history.pushState — no HTTP requests on internal nav.',
        '# But F5 sends GET /session/SESSION_001 to server → no such route → 404 without this.',
        '# Catch-all returns index.html → React boots → reconstructs view from URL.',
    ], 'SPA Catch-All Router — server.py:227')

    add_section_heading(doc, '6.2 Zustand Transient Store — 60Hz Zero-Render Architecture', 2)
    add_body(doc, 'The global store intentionally separates globalTimeMs (60Hz) from the data registry (session-level). Components subscribe transiently — no React render triggered:')
    add_zustand_flow(doc)

    add_code_block(doc, [
        '// frontend/src/components/VideoScrubber.jsx — lines 10–21',
        'const unsub = useStore.subscribe((state, prevState) => {',
        '    if (!videoRef.current || isSeekingRef.current) return',
        '    const newTime = state.globalTimeMs',
        '    if (newTime !== prevState.globalTimeMs) {',
        '        const diff = Math.abs(videoRef.current.currentTime * 1000 - newTime)',
        '        if (diff > 100) {',
        '            videoRef.current.currentTime = newTime / 1000  // Direct DOM mutation',
        '        }',
        '    }',
        '})',
        '',
        '// WHY diff > 100 guard: Prevents feedback loops.',
        '// Normal 33ms playback ticks: diff < 100ms → NO seek issued.',
        '// Chart click jumps by 500ms: diff > 100ms → seek issued.',
    ], 'VideoScrubber Transient Subscription — VideoScrubber.jsx:10')

    add_section_heading(doc, '6.3 Apache ECharts LTTB Downsampling', 2)
    add_callout(doc, 'Largest Triangle Three Buckets (LTTB) Algorithm', [
        '7,200 windows × 4 signal series = 28,800 data points WITHOUT downsampling.',
        'LTTB compresses to canvas pixel width while preserving behavioral extrema (peaks/valleys).',
        'Algorithm: for each output bucket, select the point that forms the LARGEST TRIANGLE',
        '  with the previous selected point and the centroid of the NEXT bucket.',
        'Result: Z-score anomaly spikes are ALWAYS preserved — flat segments collapse.',
        'sampling: "lttb" activated on all 4 series: L-Wrist, R-Wrist, Tremor FFT, Vocal MFCC.',
    ], C_ACCENT_PURPLE, '📊')

    add_code_block(doc, [
        '// frontend/src/components/ChartGrid.jsx — lines 60–101',
        'series: [',
        '  { name: "L-Wrist Vel", type: "line", data: leftWrist,  sampling: "lttb",',
        '    lineStyle: { color: "#3b82f6", width: 1.5 }, showSymbol: false },',
        '  { name: "R-Wrist Vel", type: "line", data: rightWrist, sampling: "lttb",',
        '    lineStyle: { color: "#10b981", width: 1.5 }, showSymbol: false },',
        '  { name: "Tremor FFT",  type: "line", data: bandPower,  sampling: "lttb",',
        '    lineStyle: { color: "#f59e0b", width: 1.5 }, showSymbol: false },',
        '  { name: "Vocal MFCC", type: "line", data: mfcc,       sampling: "lttb",',
        '    lineStyle: { color: "#8b5cf6", width: 1.5 }, showSymbol: false },',
        ']',
        '// useMemo([activeData]) — options object only recomputed on session load, NOT on 60Hz ticks',
    ], 'ECharts LTTB Series Config — ChartGrid.jsx:60')

    add_pagebreak(doc)

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX
    # ══════════════════════════════════════════════════════════════════════
    add_chapter_divider(doc, 'A',
        'Runbook & Operational Orchestration Playbook',
        '4-Terminal launch · Diagnostic troubleshooting commands',
        C_GRAY_800)

    add_section_heading(doc, 'A.1 Four-Terminal Launch Sequence', 2)

    terminals = [
        ('TERMINAL 1', '🔄 Batch Processing Daemon',
         'python app/batch_daemon.py',
         '👁️  Watchdog monitoring: .../SPOVNOB_intake/', C_ACCENT_GREEN),
        ('TERMINAL 2', '🌐 FastAPI Sidecar + Dashboard',
         'python app/server.py',
         '🚀 SPOVNOB Visualizer Sidecar running on port 8000  →  http://localhost:8000', C_ACCENT_BLUE),
        ('TERMINAL 3', '🧪 Integration Test Ingester',
         'python tests/mock_session_ingester.py',
         'Drops synthetic session_profile.json into SPOVNOB_intake/', C_ACCENT_AMBER),
        ('TERMINAL 4', '📡 SSE Ledger Stream Monitor',
         'curl -N http://localhost:8000/api/factory/stream',
         'data: {"type": "ledger_update", "data": {...}}  (pushed every 2s on change)', C_ACCENT_PURPLE),
    ]

    for num, title, cmd, expected, color in terminals:
        tbl = doc.add_table(rows=1, cols=2)
        table_no_borders(tbl)
        set_table_width(tbl, 6.5)

        badge = tbl.cell(0, 0)
        badge.width = Inches(1.2)
        set_cell_bg(badge, color)
        bp = badge.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(12)
        bp.paragraph_format.space_after = Pt(2)
        br = bp.add_run(num)
        br.font.name = 'Calibri'; br.font.size = Pt(10); br.font.bold = True; br.font.color.rgb = C_WHITE
        bs = badge.add_paragraph()
        bs.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bs.paragraph_format.space_before = Pt(0)
        bs.paragraph_format.space_after = Pt(12)
        bsr = bs.add_run(title)
        bsr.font.name = 'Calibri'; bsr.font.size = Pt(8); bsr.font.color.rgb = C_WHITE

        content = tbl.cell(0, 1)
        content.width = Inches(5.3)
        set_cell_bg(content, C_GRAY_900)
        cp = content.paragraphs[0]
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.left_indent = Pt(10)
        cr = cp.add_run(f'$ {cmd}')
        cr.font.name = 'Consolas'; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(0x86, 0xEF, 0xAC)

        ep = content.add_paragraph()
        ep.paragraph_format.space_before = Pt(2)
        ep.paragraph_format.space_after = Pt(8)
        ep.paragraph_format.left_indent = Pt(10)
        er = ep.add_run(f'  {expected}')
        er.font.name = 'Calibri'; er.font.size = Pt(8.5); er.font.color.rgb = C_GRAY_400

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_section_heading(doc, 'A.2 Diagnostic Troubleshooting Logbook', 2)
    add_data_table(doc,
        ['Symptom', 'Root Cause', 'Corrective Command'],
        [
            ['Address already in use: port 8000', 'server.py crashed, TCP socket in TIME_WAIT', 'fuser -k 8000/tcp'],
            ['CUDA not available', 'Wrong venv or CUDA drivers', 'nvidia-smi | python -c "import torch; print(torch.cuda.get_device_name(0))"'],
            ['CUDA out of memory', 'Dead subprocess leaked VRAM', 'nvidia-smi → kill -9 <PID>'],
            ['FATAL: columns missing', 'HuBERT extractor failed mid-run', 'Check metadata.json stages.acoustic_extraction.status'],
            ['SPA 404 on browser F5', 'frontend/dist/ not built', 'cd frontend && npm run build'],
            ['Ledger stuck in TENSORRT_ACTIVE', 'Subprocess killed without signal trap', 'Manually set state → INTERRUPTED via Python json patch'],
            ['IOError: EWOULDBLOCK on fcntl', 'Network copy still in progress', 'Wait for rsync/scp to complete — daemon retries on next boot'],
            ['OSError: Too many open files', 'maximize_file_descriptors() not running', 'Verify batch_daemon.py is entry point, not imported as module'],
        ], C_ACCENT_RED)

    doc.save('/Users/anshu/Documents/SPOVNOB_CLONE/SPOVNOB_SYSTEM_MASTER_MANUAL.docx')
    print('✅ DOCX generated: SPOVNOB_SYSTEM_MASTER_MANUAL.docx')

if __name__ == '__main__':
    build_doc()
